"""
Word Document (DOCX) text extractor using python-docx.
Treats the whole document as a single page.
"""
import logging
from typing import List, Dict

from docx import Document

logger = logging.getLogger(__name__)


def extract_docx_text(file_path: str) -> List[Dict]:
    """
    Extract all text from a DOCX file.
    Returns: [{page_number: 1, text: full_document_text}]
    """
    doc = Document(file_path)
    parts: List[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    full_text = "\n\n".join(parts)
    logger.info(f"DOCX extracted: {len(full_text)} chars")
    return [{"page_number": 1, "text": full_text}]
